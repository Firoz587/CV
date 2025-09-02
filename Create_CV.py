from docx import Document
from docx.shared import Pt
from datetime import datetime

# Create document
doc = Document()

# Header
doc.add_heading('MD FIROZ ISLAM', level=0)
doc.add_paragraph('Computer Science Student & Software Developer')
doc.add_paragraph('Email: mdfirozislam940@gmail.com | Phone: +8801786426404 | Location: Uttara, Dhaka-1230')

# About Me
doc.add_heading('About Me', level=1)
doc.add_paragraph(
    "I am a passionate and detail-oriented Computer Science student at Northern University Bangladesh, "
    "with a keen interest in software development, database management, and problem solving. "
    "I have honed my programming and algorithmic skills through coding competitions and enjoy sharing knowledge."
)

# Education
doc.add_heading('Education', level=1)
education = [
    ('BSC in Computer Science', '2022 – Present', 'Northern University Bangladesh'),
    ('HSC', '2019 – 2021', 'Armed Police Battalion School & College, Bogura (Result: Golden A+)'),
    ('SSC', '2012 – 2019', 'Dighirhat High School (Result: Golden A+)')
]
for title, date, school in education:
    doc.add_heading(title, level=2)
    doc.add_paragraph(f'{date}\n{school}')

# Projects
doc.add_heading('Projects', level=1)
projects = [
    ('Bank Management System', 'Python, File Handling',
     'Console-based app to manage banking operations including account creation, transactions, and authentication.'),
    ('E-Commerce Project (Car Shop)', 'Django, HTML, CSS, JavaScript',
     'Online car shop with product listings, authentication, and shopping cart functionality.'),
    ('Smart-Care & Hotel Management System (API)', 'Django REST Framework, PostgreSQL',
     'RESTful APIs for smart-care and hotel management systems with full CRUD functionality.')
]
for title, tech, desc in projects:
    doc.add_heading(title, level=2)
    doc.add_paragraph(f'Technologies: {tech}')
    doc.add_paragraph(desc)

# Skills
doc.add_heading('Skills', level=1)
skills = ['C/C++', 'Python', 'Django', 'SQL', 'MySQL', 'PostgreSQL', 'HTML', 'CSS', 'Bootstrap', 'Tailwind', 'JavaScript', 'Problem Solving']
doc.add_paragraph(', '.join(skills))

# Achievements
doc.add_heading('Achievements', level=1)
achievements = [
    'Codeforces Rating: Pupil (Max Rating 1238) - https://codeforces.com/profile/Firoz_587',
    'Codechef Rating: 3 Star (Max Rating 1615) - https://www.codechef.com/users/firoz_587',
    'Python Certificate (CGPA-4)'
]
for ach in achievements:
    doc.add_paragraph(ach, style='List Bullet')

# Languages
doc.add_heading('Languages', level=1)
languages = [('Bengali', 'Native'), ('English', 'Fluent')]
for lang, level in languages:
    doc.add_paragraph(f'{lang} - {level}')

# Hobbies & Interests
doc.add_heading('Interests & Hobbies', level=1)
hobbies = ['Coding & Problem Solving', 'Programming', 'Reading']
for hobby in hobbies:
    doc.add_paragraph(hobby, style='List Bullet')

# Footer / Last Updated
doc.add_paragraph(f'\nLast updated: {datetime.now().strftime("%B %Y")}')

# Save
doc.save('MD_Firoz_Islam_CV.docx')
print("DOCX file created successfully!")
